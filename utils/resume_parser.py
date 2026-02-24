"""
Resume / cover letter parser — converts uploaded files to plain text.

Supported formats:
  PDF   → pypdf  (pip install pypdf)
  DOCX  → python-docx  (pip install python-docx)
  TXT   → built-in (no extra package)

Usage:
    text = parse_uploaded_file(st_uploaded_file)
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_uploaded_file(uploaded_file) -> str:
    """
    Parse a Streamlit UploadedFile object into plain text.

    Args:
        uploaded_file: streamlit.runtime.uploaded_file_manager.UploadedFile

    Returns:
        Extracted plain text string (may be empty if parsing failed).

    Raises:
        ValueError: Unsupported file extension.
        ImportError: Required parsing package not installed.
        Exception: Underlying parsing error.
    """
    name = uploaded_file.name.lower()
    raw_bytes = uploaded_file.read()

    if name.endswith(".pdf"):
        return _parse_pdf(raw_bytes)
    elif name.endswith(".docx"):
        return _parse_docx(raw_bytes)
    elif name.endswith((".txt", ".rtf", ".md")):
        # Try UTF-8, fall back to latin-1
        try:
            return raw_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            return raw_bytes.decode("latin-1", errors="replace").strip()
    else:
        raise ValueError(
            f"Unsupported file type '{uploaded_file.name}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        import pypdf
    except ImportError:
        raise ImportError(
            "pypdf is not installed.\n"
            "Run:  C:\\Temp\\ClaudeCode\\python313\\Scripts\\pip.exe install pypdf"
        )

    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        result = "\n".join(pages_text).strip()
        if not result:
            logger.warning("PDF parsed but no text extracted — may be a scanned/image PDF.")
        return result
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise RuntimeError(f"Could not read PDF: {e}") from e


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is not installed.\n"
            "Run:  C:\\Temp\\ClaudeCode\\python313\\Scripts\\pip.exe install python-docx"
        )

    try:
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise RuntimeError(f"Could not read DOCX: {e}") from e


def save_upload_text(text: str, name: str, uploads_dir: Path) -> None:
    """
    Save extracted text to disk for persistence across Streamlit restarts.

    Args:
        text: Extracted text content.
        name: Base name for the file ('resume' or 'cover_letter').
        uploads_dir: Directory to save into (created if needed).
    """
    uploads_dir.mkdir(parents=True, exist_ok=True)
    dest = uploads_dir / f"{name}.txt"
    dest.write_text(text, encoding="utf-8")
    logger.info(f"Saved {name} text to {dest} ({len(text):,} chars)")


def load_saved_text(name: str, uploads_dir: Path) -> str:
    """
    Load previously saved text from disk.

    Returns empty string if no file found.
    """
    dest = uploads_dir / f"{name}.txt"
    if dest.exists():
        try:
            return dest.read_text(encoding="utf-8")
        except Exception as e:
            logger.warning(f"Could not read saved {name}: {e}")
    return ""
